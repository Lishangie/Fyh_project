import os
import argparse
from typing import Optional
try:
    # prefer the installed langgraph package if available
    from langgraph import StateGraph, END
    from langgraph.checkpoint import SqliteSaver
except Exception:
    # fallback to local emulator
    from langgraph.graph import StateGraph, END
    from langgraph.checkpoint.sqlite import SqliteSaver

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm

from state import ReportState
from agents.researcher import research_node
from agents.writer import writer_node
from agents.coder import coder_visualizer_node
from agents.error_resolver import error_resolver_node
from agents.memory import feedback_processor_node

def document_assembler_node(state: ReportState) -> dict:
    # Avoid re-running assembly if already completed in this session
    if state.get("assembled_done"):
        return {"execution_errors": []}

    print("--- Запуск генерации эталонного DOCX файла ---")
    template_path = os.path.join("assets", "template_gost.docx")
    os.makedirs("assets", exist_ok=True)
    os.makedirs("artifacts", exist_ok=True)

    # Если шаблон отсутствует — создаём минимальный шаблон с Jinja2-тегами
    if not os.path.exists(template_path):
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("{{ report_title }}")
            doc.add_paragraph("")
            doc.add_paragraph("{{ main_body_text }}")
            doc.add_paragraph("{{ dynamic_img_0 }}")
            doc.add_paragraph("{{ dynamic_img_1 }}")
            doc.save(template_path)
            print(f"Создан базовый шаблон: {template_path}")
        except Exception as e:
            return {"execution_errors": [f"Не удалось создать шаблон: {e}"]}

    try:
        doc = DocxTemplate(template_path)
    except Exception as e:
        return {"execution_errors": [f"Ошибка загрузки шаблона docx: {e}"]}

    # Validate artifact paths exist before assembly
    artifact_paths = list(state.get("artifact_paths", []))
    missing = [p for p in artifact_paths if not os.path.exists(p)]
    if missing:
        errs = list(state.get("execution_errors", []))
        for m in missing:
            errs.append(f"assembler: missing artifact {m}")
        return {"execution_errors": errs}

    images_context = {}
    for idx, path in enumerate(artifact_paths):
        if os.path.exists(path):
            images_context[f"dynamic_img_{idx}"] = InlineImage(doc, path, width=Mm(160))
        else:
            print(f"Внимание: изображение {path} не найдено.")

    context = {
        "report_title": state.get("task_description", "Академический Отчет"),
        "main_body_text": state.get("draft_text", ""),
        "data_tables": state.get("dynamic_tables", []),
        **images_context
    }

    try:
        doc.render(context)
        output_file = os.path.join("artifacts", "Final_Academic_Report.docx")
        doc.save(output_file)
        print(f"--- Документ успешно скомпилирован: {output_file} ---")
        return {"execution_errors": [], "assembled_done": True}
    except Exception as e:
        return {"execution_errors": [str(e)]}

def route_code_execution(state: ReportState):
    errors = state.get("execution_errors", [])
    return "coder_error" if errors else "coder_ok"


def writer_router(state: ReportState):
    errors = state.get("execution_errors", [])
    return "writer_error" if errors else "writer_ok"


def error_resolution_router(state: ReportState):
    # Decide where to go after error_resolver_node sets last_resolution and last_failed_node
    lr = state.get("last_resolution")
    failed = state.get("last_failed_node")
    if lr == "retry":
        # route back to the failing node
        if failed and "writer" in failed:
            return "writer_node"
        if failed and "coder" in failed:
            return "coder_node"
        return "coder_node"
    return "assembler_node"

def build_autonomous_graph():
    workflow = StateGraph(ReportState)
    workflow.add_node("researcher_node", research_node)
    workflow.add_node("writer_node", writer_node)
    workflow.add_node("coder_node", coder_visualizer_node)
    workflow.add_node("error_resolver_node", error_resolver_node)
    workflow.add_node("feedback_processor_node", feedback_processor_node)
    workflow.add_node("assembler_node", document_assembler_node)

    workflow.set_entry_point("researcher_node")
    workflow.add_edge("researcher_node", "writer_node")

    # Writer -> either coder or error resolver
    workflow.add_conditional_edges(
        "writer_node",
        writer_router,
        {
            "writer_ok": "coder_node",
            "writer_error": "error_resolver_node",
        },
    )

    # Coder -> assembler or error resolver
    workflow.add_conditional_edges(
        "coder_node",
        route_code_execution,
        {
            "coder_ok": "assembler_node",
            "coder_error": "error_resolver_node",
        },
    )

    # After resolving an error, route according to resolver's decision
    workflow.add_conditional_edges(
        "error_resolver_node",
        error_resolution_router,
        {
            "writer_node": "writer_node",
            "coder_node": "coder_node",
            "assembler_node": "assembler_node",
        },
    )

    # Feedback processor should reuse the same resolution router to decide where to retry
    workflow.add_conditional_edges(
        "feedback_processor_node",
        error_resolution_router,
        {
            "writer_node": "writer_node",
            "coder_node": "coder_node",
            "assembler_node": "assembler_node",
        },
    )

    # Assembly may fail if artifacts are missing; route accordingly
    def assembler_router(state: ReportState):
        errors = state.get("execution_errors", [])
        human_fb = (state.get("human_feedback") or "").strip()
        if errors:
            return "assembler_error"
        if human_fb:
            return "assembler_feedback"
        return "assembler_ok"

    workflow.add_conditional_edges(
        "assembler_node",
        assembler_router,
        {
            "assembler_ok": END,
            "assembler_error": "error_resolver_node",
            "assembler_feedback": "feedback_processor_node",
        },
    )

    memory = SqliteSaver.from_conn_string("sqlite_checkpoints.db")
    # Allow tests to auto-approve HITL assembly by disabling the interrupt
    auto = os.environ.get("AUTO_APPROVE", "0") in ("1", "true", "True")
    interrupts = [] if auto else ["assembler_node"]
    graph = workflow.compile(checkpointer=memory, interrupt_before=interrupts)
    return graph


def run_cli(argv: Optional[list] = None):
    parser = argparse.ArgumentParser()
    parser.add_argument("--thread-id", default="mirea_report_session_01", help="ID of the graph thread/session")
    parser.add_argument("-y", "--yes", action="store_true", help="auto-approve HITL checkpoint and continue")
    args = parser.parse_args(argv)

    app_graph = build_autonomous_graph()
    thread_config = {"configurable": {"thread_id": args.thread_id}}
    initial_state = {
        "task_description": "Проектирование и анализ отказоустойчивой IoT системы",
        "context_data": "Логи нагрузочного тестирования sensors_data.csv",
        "artifact_paths": [],
        "dynamic_tables": [],
        "execution_errors": []
    }

    print(">>> Инициализация графа генерации...")
    for event in app_graph.stream(initial_state, config=thread_config):
        for node_name, node_state in event.items():
            print(f"=== Завершено выполнение узла: {node_name} ===")

    # Inspect checkpoint to see whether we paused before a HITL node
    state_snapshot = app_graph.get_state(thread_config)
    current_state = state_snapshot.values
    current_node = state_snapshot.current_node

    if current_node is None:
        print("\n>>> Граф завершён — нет чекпоинта. Проверьте artifacts/ for output.")
        return

    print("\n>>> Процесс приостановлен фреймворком (Human-in-the-Loop). Откройте artifacts/ и assets/ для проверки.")
    print(f"Сгенерировано изображений: {len(current_state.get('artifact_paths', []))}")
    print(f"Сгенерировано таблиц: {len(current_state.get('dynamic_tables', []))}")

    auto = args.yes or os.environ.get("AUTO_APPROVE", "0") in ("1", "true", "True")
    if auto:
        approve = True
        feedback_input = None
        print("Авто-одобрение включено — продолжаем сборка.")
    else:
        user_input = input(
            "\nОдобрить промежуточные результаты и запустить компиляцию DOCX?\n"
            "Введите 'y' для подтверждения, 'n' для отмены, или введите текстовый отзыв для обучения: "
        )
        ui = (user_input or "").strip()
        if ui.lower() == 'y':
            approve = True
            feedback_input = None
        elif ui.lower() == 'n' or ui == '':
            approve = False
            feedback_input = None
        else:
            # treat any other non-empty input as feedback to process
            approve = True
            feedback_input = ui

    if approve:
        # If user provided freeform feedback, save it into the checkpoint so
        # the graph can route to the feedback processor when resumed.
        if feedback_input:
            thread_id = thread_config.get("configurable", {}).get("thread_id", "default_thread")
            current_state["human_feedback"] = feedback_input
            # save updated checkpoint at the same paused node
            try:
                app_graph.checkpointer.save(thread_id, current_state, current_node)
            except Exception as e:
                print(f"Не удалось сохранить отзыв в чекпоинт: {e}")

        print(">>> Возобновление выполнения графа...")
        for event in app_graph.stream(None, config=thread_config):
            for node_name, node_state in event.items():
                print(f"=== Завершено выполнение узла: {node_name} ===")
    else:
        print(">>> Сборка остановлена. Обновите state (human_feedback) для повторной генерации.")


if __name__ == "__main__":
    run_cli()
